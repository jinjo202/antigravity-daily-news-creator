import os
import re
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def style_run(run, font_name='바탕체', size_pt=15, bold=False, underline=False, color_rgb=None, is_footnote=False):
    """Applies strict typography to a run in Batang font with explicit color control."""
    actual_size = 11 if is_footnote else size_pt
    actual_color = RGBColor(0, 0, 255) if is_footnote else (color_rgb if color_rgb is not None else RGBColor(0, 0, 0))
    actual_bold = False if is_footnote else bold
    actual_underline = False if is_footnote else underline
    
    run.font.name = font_name
    run.font.size = Pt(actual_size)
    run.bold = actual_bold
    run.underline = actual_underline
    run.font.color.rgb = actual_color
    
    # XML fix for East Asian font rendering in MS Word
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph(paragraph, space_before_pt, space_after_pt, left_indent_pt, first_line_indent_pt, word_wrap_off=True):
    """Formats paragraph spacing, indents, and line spacing according to guidelines."""
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before_pt)
    p_format.space_after = Pt(space_after_pt)
    p_format.left_indent = Pt(left_indent_pt)
    p_format.first_line_indent = Pt(first_line_indent_pt)
    p_format.line_spacing = 1.3
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Enforce word wrap to prevent Hangul word splitting (한글 단어 잘림 방지)
    pPr = paragraph._p.get_or_add_pPr()
    
    # 금칙 처리 활성화
    kinsoku = OxmlElement('w:kinsoku')
    kinsoku.set(qn('w:val'), '1')
    pPr.append(kinsoku)
    
    # 단어 단위 줄 바꿈 활성화
    wordWrap = OxmlElement('w:wordWrap')
    wordWrap.set(qn('w:val'), '1')
    pPr.append(wordWrap)

def convert_to_nominal(sentence):
    """Converts a polite Korean sentence to nominal-ending (개조식) form."""
    sentence = sentence.strip()
    if not sentence:
        return ""
    
    clean_sent = sentence.replace("**", "").replace("__", "")
    has_period = clean_sent.endswith('.')
    if has_period:
        clean_sent = clean_sent[:-1].strip()
        
    replacements = [
        (r'하락 했습니다$', '하락함'),
        (r'하락했습니다$', '하락함'),
        (r'상승했습니다$', '상승함'),
        (r'급등했습니다$', '급등함'),
        (r'급락했습니다$', '급락함'),
        (r'기록했습니다$', '기록함'),
        (r'동향입니다$', '동향임'),
        (r'변수입니다$', '변수임'),
        (r'기대하고 있습니다$', '기대하고 있음'),
        (r'기대하고있습니다$', '기대하고 있음'),
        (r'생각됩니다$', '생각됨'),
        (r'보였습니다$', '보임'),
        (r'전망됩니다$', '전망됨'),
        (r'출발했지만$', '출발하였으나'),
        (r'하락했고$', '하락함'),
        (r'상승했고$', '상승함'),
        (r'내려갔으나$', '하락하였으나'),
        (r'제시됐기 때문에$', '제시됨에 따라'),
        (r'제시되었기 때문에$', '제시됨에 따라'),
        (r'있을지가 변수입니다$', '있을지가 변수임'),
        (r'기대하고 있습니다$', '기대함'),
        (r'생각됩니다$', '판단됨'),
        (r'했습니다$', '함'),
        (r'하였습니다$', '함'),
    ]
    
    for pattern, repl in replacements:
        clean_sent = re.sub(pattern, repl, clean_sent)
        
    if has_period:
        clean_sent += '.'
    return clean_sent

def parse_us_report(text):
    """Parses raw US market report text into structured sections."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    
    title = ""
    greeting = ""
    sec1_lines = []
    ytd_line = ""
    sec2_lines = []
    semi_details = ""
    bond_line = ""
    sec3_lines = []
    sec4_lines = []
    
    state = "TITLE"
    for line in lines:
        clean_line = line.replace("**", "").replace("__", "")
        if clean_line.lower().startswith("title"):
            title = line
            state = "GREETING"
            continue
        if "안녕하십니까" in clean_line:
            greeting = line
            state = "SEC1"
            continue
            
        if clean_line.startswith("*") and "연초대비" in clean_line:
            ytd_line = line
            state = "SEC2"
            continue
            
        if state == "SEC1":
            sec1_lines.append(line)
        elif state == "SEC2":
            if clean_line.startswith("(") and "반도체" in clean_line:
                semi_details = line
            elif "10년 국채 금리" in clean_line or "국채 금리" in clean_line:
                bond_line = line
                state = "SEC3"
            else:
                sec2_lines.append(line)
        elif state == "SEC3":
            if clean_line.startswith("한편") or "국내 증시" in clean_line:
                sec4_lines.append(line)
                state = "SEC4"
            else:
                sec3_lines.append(line)
        elif state == "SEC4":
            if "감사합니다" not in clean_line:
                sec4_lines.append(line)
                
    # Extract date from Title
    date_str = datetime.now().strftime("%Y년 %m월 %d일")
    if title:
        colon_idx = title.find(':')
        if colon_idx != -1:
            title_val = title[colon_idx + 1:].strip()
            match = re.search(r'\((\d+)/(\d+)\)', title_val)
            if match:
                month = match.group(1)
                day = match.group(2)
                date_str = f"2026년 {month}월 {day}일"
                
    return {
        'date': date_str,
        'sec1_text': " ".join(sec1_lines),
        'ytd_line': ytd_line,
        'sec2_text': " ".join(sec2_lines),
        'semi_details': semi_details,
        'bond_line': bond_line,
        'sec3_text': "\n".join(sec3_lines),
        'sec4_text': "\n".join(sec4_lines)
    }

def wrap_korean_text(text, max_len=30):
    """Wraps Korean text strictly by spaces, ensuring no line exceeds max_len characters.
    Words are never split across lines. Certain price/index/YTD lines are kept on a single line."""
    line_stripped = text.strip()
    
    # Check if this line should not be wrapped (YTD, Samsung stocks, exchange rates, bond yields)
    if line_stripped.startswith('※') or (line_stripped.startswith('*') and '연초대비' in line_stripped):
        return [text]
    if line_stripped.startswith('(') and any(kw in line_stripped for kw in ['전자', '화재', '생명']):
        return [text]
    if ('△' in line_stripped or '+' in line_stripped or '%' in line_stripped) and any(kw in line_stripped for kw in ['원', 'bp', '%', 'pt', '지수', '상승', '하락', '금리']):
        if len(line_stripped) < 75:
            return [text]
            
    is_bold_header = text.startswith('**') and text.endswith('**')
    if is_bold_header:
        text = text[2:-2].strip()
        
    words = text.split()
    lines = []
    current_line = []
    current_len = 0
    
    for word in words:
        visible_word = word.replace("**", "").replace("__", "")
        word_len = len(visible_word)
        
        extra = 1 if current_line else 0
        if current_len + word_len + extra > max_len:
            if current_line:
                lines.append(" ".join(current_line))
                current_line = [word]
                current_len = word_len
            else:
                current_line = [word]
                current_len = word_len
        else:
            current_line.append(word)
            current_len += word_len + extra
            
    if current_line:
        lines.append(" ".join(current_line))
        
    result = [l for l in lines if l]
    if is_bold_header:
        result = [f"**{l}**" for l in result]
    return result

def create_us_report_document(data, output_path):
    """Generates the daily US market report Word document matching the sample style exactly."""
    doc = Document()
    
    # 1. Page margins (상하좌우 여백 설정: 상 1.18인치 = 85.05pt, 하/좌/우 1.0인치 = 72pt)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Pt(85.05)
        section.bottom_margin = Pt(72.0)
        section.left_margin = Pt(72.0)
        section.right_margin = Pt(72.0)
        
    # Get raw report text
    raw_text = data if isinstance(data, str) else data.get('raw_text', '')
    if not raw_text and isinstance(data, dict):
        # Fallback: reconstruct raw text from sections if raw_text key is missing
        parts = []
        sec1 = data.get('sec1_text', '')
        if sec1: parts.append(sec1)
        sec2 = data.get('sec2_text', '')
        if sec2: parts.append(sec2)
        semi = data.get('semi_details', '')
        if semi: parts.append(semi)
        bond = data.get('bond_line', '')
        if bond: parts.append(bond)
        sec3 = data.get('sec3_text', '')
        if sec3: parts.append(sec3)
        sec4 = data.get('sec4_text', '')
        if sec4: parts.append(sec4)
        raw_text = "\n\n".join(parts)

    # Process line by line
    lines = raw_text.splitlines()
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip title line at the top of the document (sample does not show it)
        line_clean = line_stripped.replace('*', '').strip()
        if line_clean.lower().startswith('title') or line_clean.lower().startswith('**title'):
            continue
        if line_stripped.startswith('**') and line_stripped.endswith('**') and '시황(' in line_stripped:
            continue
            
        # Empty line maps to empty paragraph (preserves spacing blocks)
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Wrap the line to max 30 characters
        wrapped_lines = wrap_korean_text(line_stripped, max_len=30)
        for w_line in wrapped_lines:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(8)
            p_format.line_spacing = 1.15
            
            # Enable kinsoku and wordWrap for Korean (한글 단어 잘림 방지)
            pPr = p._p.get_or_add_pPr()
            kinsoku = OxmlElement('w:kinsoku')
            kinsoku.set(qn('w:val'), '1')
            pPr.append(kinsoku)
            
            wordWrap = OxmlElement('w:wordWrap')
            wordWrap.set(qn('w:val'), '1')
            pPr.append(wordWrap)
            
            # Check if YTD line (starts with ※ or *)
            is_ytd = w_line.startswith('※') or (w_line.startswith('*') and '연초대비' in w_line)
            color = RGBColor(0, 0, 255) if is_ytd else RGBColor(0, 0, 0)
            
            # Parse markdown bold (**...**)
            parts = re.split(r'(\*\*[^*]+\*\*)', w_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run_text = part[2:-2]
                    bold = True
                else:
                    run_text = part
                    bold = False
                    
                if run_text:
                    run = p.add_run(run_text)
                    run.font.name = '바탕체'
                    run.font.size = Pt(11)
                    run.bold = bold
                    run.font.color.rgb = color
                    
                    # XML fix for East Asian fonts rendering
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), '바탕체')
                    rFonts.set(qn('w:hAnsi'), '바탕체')
                    rFonts.set(qn('w:eastAsia'), '바탕체')
                
    doc.save(output_path)
    print(f"[성공] 미국 시황 보고서 워드 문서 생성 완료: {output_path}")

if __name__ == '__main__':
    # Test generation
    sample_text = """
    Title : 일일 금융시장 동향(6/16)
    안녕하십니까
    6월 16일 국내외 금융시장 동향입니다.
    미국 증시는 미/이란 종전 합의로 유가 추가 하락하며 상승 출발했지만, FOMC를 앞둔 경계감에 하락했고, 유럽 증시는 독일 경기기대지수 호조 영향에 상승했습니다.
    * 6/16(연초대비): S&P500 △0.6%(+9.7), 나스닥 △1.2(+13.5), Stoxx50 +0.4(+8.1)
    국제 유가는 △5% 하락한 $76 수준까지 내려갔으나, FOMC 대기 심리에 반도체주 위주로 차익실현이 발생하며 미 증시가 하락했습니다.
    (미 반도체 지수 △6%, 마이크론 △6%, 인텔 △9% 등).
    미 10년 국채 금리는 유가 하락 영향에 △4bp(4.44%) 하락 했습니다.
    6월 FOMC에서는 기준금리 동결이 유력하나, 3월 FOMC 당시 점도표는 연내 1회 '인하'로 제시됐기 때문에, 이번에 점도표가 연내 1회 '인상'으로 변경되거나, 연내 동결로 제시되면서도 케빈 워시 의장의 매파적 기자회견이 있을지가 변수입니다.
    단, 시장에서 이미 연내 금리 1회 인상을 반영해온 가운데, 최근 유가 하락으로 당초 예상보다는 덜 매파적인 기자회견이 나올 것으로 기대하고 있습니다.
    한편, 국내 증시에서 외국인은 최근 3일 4.8조원을 순매수하며 수급 변화 조짐이 보였습니다. (5.7~6.11까지 △76조원 순매도)
    외국인 순매수가 유가가 본격 하락하면서 유가 피해국으로 여겨진 신흥국에 대한 매도가 완화되며 나타난 것이라면 추가 순매수가 진행될 수 있기 때문에 증시에 긍정적인 요인이 될 것으로 생각됩니다.
    감사합니다.
    """
    data = parse_us_report(sample_text)
    create_us_report_document(data, "test_us_report.docx")
