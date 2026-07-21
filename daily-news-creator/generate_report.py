import os
import re
import sys
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==============================================================================
# DOCX XML STYLING HELPERS
# ==============================================================================
def style_run(run, font_name='바탕체', size_pt=15, bold=False, underline=False, color_rgb=None, is_footnote=False):
    """Applies strict typography to a run in Batang font with explicit color control."""
    actual_size = 11 if is_footnote else size_pt
    actual_color = RGBColor(0, 0, 255) if is_footnote else (color_rgb if color_rgb is not None else RGBColor(0, 0, 0))
    actual_bold = False if is_footnote else bold
    actual_underline = False if is_footnote else underline
    
    run.font.name = font_name
    run.font.size = Pt(actual_size)
    run.bold = actual_bold
    run.underline = actual_underline
    run.font.color.rgb = actual_color
    
    # XML fix for East Asian font rendering in MS Word
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph(paragraph, space_before_pt, space_after_pt, left_indent_pt, first_line_indent_pt, word_wrap_off=True):
    """Formats paragraph spacing, indents, and line spacing according to guidelines (default word_wrap_off=True)."""
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before_pt)
    p_format.space_after = Pt(space_after_pt)
    p_format.left_indent = Pt(left_indent_pt)
    p_format.first_line_indent = Pt(first_line_indent_pt)
    p_format.line_spacing = 1.3  # Fixed line spacing 1.3
    
    # Use JUSTIFY alignment for Korean public documents to align both margins like a rectangle.
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Enforce word wrap to prevent Hangul word splitting (한글 단어 잘림 방지)
    pPr = paragraph._p.get_or_add_pPr()
    
    # 금칙 처리 활성화
    kinsoku = OxmlElement('w:kinsoku')
    kinsoku.set(qn('w:val'), '1')
    pPr.append(kinsoku)
    
    # 단어 단위 줄 바꿈 활성화
    wordWrap = OxmlElement('w:wordWrap')
    wordWrap.set(qn('w:val'), '1')
    pPr.append(wordWrap)

# ==============================================================================
# NOMINAL (GAEJOSIK) ENDING CONVERTER
# ==============================================================================
def convert_to_nominal(sentence):
    """Converts a polite Korean sentence to nominal-ending (개조식) form."""
    sentence = sentence.strip().replace("**", "").replace("__", "")
    if not sentence:
        return ""
        
    has_period = sentence.endswith('.')
    if has_period:
        sentence = sentence[:-1].strip()
        
    # Replacements for typical Korean polite verb endings to nominal form
    replacements = [
        (r'경신했습니다$', '경신함'),
        (r'경신하였습니다$', '경신함'),
        (r'기록했습니다$', '기록함'),
        (r'기록하였습니다$', '기록함'),
        (r'폭락했습니다$', '폭락함'),
        (r'하락했습니다$', '하락함'),
        (r'급락했습니다$', '급락함'),
        (r'상승했습니다$', '상승함'),
        (r'급등했습니다$', '급등함'),
        (r'지속했습니다$', '지속함'),
        (r'지속하였습니다$', '지속함'),
        (r'보였습니다$', '보임'),
        (r'전망입니다$', '전망됨'),
        (r'전망됩니다$', '전망됨'),
        (r'있습니다$', '있음'),
        (r'없습니다$', '없음'),
        (r'합니다$', '함'),
        (r'드립니다$', '드림'),
        (r'유지했으나$', '유지하였으나'),
        (r'하락하였으나$', '하락하였으나'),
        (r'상승하였으나$', '상승하였으나'),
        (r'했습니다$', '함'),
        (r'하였습니다$', '함'),
        (r'되었다$', '됨'),
        (r'되었습니다$', '됨'),
        (r'일으켜$', '일으킴'),
        (r'지속했고$', '지속함'),
        (r'상승했고$', '상승함'),
        (r'하락했고$', '하락함'),
    ]
    
    for pattern, repl in replacements:
        sentence = re.sub(pattern, repl, sentence)
        
    if has_period:
        sentence += '.'
    return sentence

# ==============================================================================
# TEXT PARSING ENGINE
# ==============================================================================
def parse_raw_text(raw_text):
    """Intelligently parses the raw daily report text into structured dict data."""
    lines = [line.strip() for line in raw_text.splitlines()]
    
    title = ""
    greeting = ""
    asian_lines = []
    ytd_line = ""
    korean_lines = []
    stock_details_line = ""
    fx_bond_lines = []
    outlook_lines = []
    
    state = "GREETING"
    
    for line in lines:
        if not line:
            continue
        
        # Clean line by removing markdown formatting (** and __) and leading bullet symbols
        clean_line = line.strip().replace("**", "").replace("__", "")
        # Remove leading symbols like *, _, -, tab, space, bullet points
        clean_line = re.sub(r'^[\s*_~•\-▶▷■□※]*', '', clean_line).strip()
        
        if line.lower().startswith("title") or clean_line.lower().startswith("title"):
            title = line
            continue
            
        if "보고 드립니다" in clean_line or "시황 보고" in clean_line or "보고드립니다" in clean_line:
            greeting = line
            state = "ASIAN_MARKET"
            continue
            
        if (line.startswith("※") or clean_line.startswith("※")) and "연초대비" in clean_line:
            ytd_line = line
            continue
            
        # State transitions based on cleaned line
        if state == "ASIAN_MARKET":
            if (clean_line.startswith("한국 증시") or 
                clean_line.startswith("코스피는") or 
                clean_line.startswith("코스피") or 
                "코스피는" in clean_line or 
                "한국 증시" in clean_line or
                "코스닥" in clean_line):
                state = "KOREAN_MARKET"
                
        if state == "KOREAN_MARKET":
            if (clean_line.startswith("원달러") or 
                clean_line.startswith("원/달러") or 
                clean_line.startswith("환율은") or 
                clean_line.startswith("환율") or
                clean_line.startswith("국고채") or
                clean_line.startswith("금리")):
                state = "FX_BONDS"
                
        if state == "FX_BONDS":
            if (clean_line.startswith("향후에는") or 
                clean_line.startswith("향후") or 
                clean_line.startswith("금주는") or 
                clean_line.startswith("이번 주") or 
                clean_line.startswith("이번주") or
                clean_line.startswith("글로벌") or
                clean_line.startswith("주요")):
                state = "OUTLOOK"
                
        # Append lines (using original lines to preserve formatting like ** and __)
        if state == "ASIAN_MARKET":
            asian_lines.append(line)
        elif state == "KOREAN_MARKET":
            if clean_line.startswith("(") and "전자" in clean_line:
                stock_details_line = line
            else:
                korean_lines.append(line)
        elif state == "FX_BONDS":
            fx_bond_lines.append(line)
        elif state == "OUTLOOK":
            if "감사합니다" not in clean_line:
                outlook_lines.append(line)
                
    # Build data dictionary
    data = {
        'raw_text': raw_text,
        'date': datetime.now().strftime("%Y년 %m월 %d일"),
        'section1_text': "\n".join(asian_lines),
        'asian_indices': [],
        'section2_text1': "",
        'section2_text2': "",
        'stocks': [],
        'section3_text': "\n".join(fx_bond_lines),
        'macro_indicators': [],
        'quote': {},
        'section4_text': "\n".join(outlook_lines),
        'events': []
    }
    
    # Parse Title for Date
    if title:
        colon_idx = title.find(':')
        if colon_idx != -1:
            title_val = title[colon_idx + 1:].strip()
            match = re.search(r'\((\d+)/(\d+)\)', title_val)
            if match:
                month = match.group(1)
                day = match.group(2)
                data['date'] = f"2026년 {month}월 {day}일"
                
    # Parse YTD indices
    if ytd_line:
        colon_idx = ytd_line.find(':')
        if colon_idx != -1:
            entries_part = ytd_line[colon_idx + 1:].strip()
            entries = [e.strip() for e in entries_part.split(',')]
            for entry in entries:
                m_country = re.match(r'^([가-힣]+)\s+(.*)$', entry)
                if m_country:
                    country = m_country.group(1)
                    rest = m_country.group(2).strip()
                    c_name = f"{country} KOSPI" if country == "한국" else (f"{country} Nikkei 225" if country == "일본" else f"{country} Shanghai")
                    
                    if "휴장" in rest:
                        data['asian_indices'].append({
                            'country': c_name,
                            'change': '휴장',
                            'ytd': '휴장'
                        })
                    else:
                        m_val = re.match(r'^([+\-△▲▼▽]?[0-9.]+[^()]*)\s*\(\s*([+\-△▲▼▽]?[0-9.]+[^()]*)\s*\)', rest)
                        if m_val:
                            change = m_val.group(1).strip()
                            ytd = m_val.group(2).strip()
                            if not change.endswith('%') and any(char.isdigit() for char in change):
                                change += '%'
                            if not ytd.endswith('%') and any(char.isdigit() for char in ytd):
                                ytd += '%'
                            data['asian_indices'].append({
                                'country': c_name,
                                'change': change,
                                'ytd': ytd
                            })
            
    # Parse Korean Market text (split by foreigners paragraph)
    section2_text1_list = []
    section2_text2_list = []
    in_part2 = False
    for p in korean_lines:
        if p.startswith("외국인"):
            in_part2 = True
        if in_part2:
            section2_text2_list.append(p)
        else:
            section2_text1_list.append(p)
            
    data['section2_text1'] = "\n".join(section2_text1_list)
    data['section2_text2'] = "\n".join(section2_text2_list)
    
    # Parse stocks
    if stock_details_line:
        matches = re.findall(r'([가-힣a-zA-Z0-9]+)\s+([+\-△▲▼▽]?[0-9.]+%?)', stock_details_line)
        for name, change in matches:
            full_name = f"삼성{name}" if name in ['전자', '화재', '생명'] else name
            if not change.endswith('%') and any(char.isdigit() for char in change):
                change += '%'
            data['stocks'].append({
                'name': full_name,
                'change': change
            })
            
    # Parse macro indicators
    s3_text = data['section3_text']
    
    # USD/KRW Exchange Rate
    usd_match = re.search(r'([+\-△▲▼▽]?[0-9.]+)\s*원\s*\(\s*([\d,.]+)\s*원\s*\)', s3_text)
    if usd_match:
        data['macro_indicators'].append({
            'name': '원/달러 환율 (USD/KRW)',
            'value': usd_match.group(2) + "원",
            'change': usd_match.group(1) + "원"
        })
    else:
        usd_match_nat = re.search(r'([+\-△▲▼▽]?[0-9.]+)\s*원\s*(?:내린|상승한|하락한|오른)?\s*([\d,.]+)\s*원', s3_text)
        if usd_match_nat:
            # Find the specific line containing this match
            matched_line = ""
            for line in s3_text.splitlines():
                if usd_match_nat.group(0) in line:
                    matched_line = line
                    break
            change_val = usd_match_nat.group(1)
            change_str = change_val + "원"
            if "내린" in matched_line or "하락한" in matched_line:
                if not any(char in change_str for char in ['△', '-', '▼']):
                    change_str = "△" + change_str
            data['macro_indicators'].append({
                'name': '원/달러 환율 (USD/KRW)',
                'value': usd_match_nat.group(2) + "원",
                'change': change_str
            })
            
    # Bond Yield
    bond_match = re.search(r'([+\-△▲▼▽]?[0-9.]+)\s*bp\s*\(\s*([\d,.]+)\s*%\s*\)', s3_text)
    if bond_match:
        data['macro_indicators'].append({
            'name': '국고채 금리 (10년물)',
            'value': bond_match.group(2) + "%",
            'change': bond_match.group(1) + "bp"
        })
    else:
        bond_match_nat = re.search(r'([+\-△▲▼▽]?[0-9.]+)\s*bp\s*(?:상승한|하락한|내린|오른)?\s*([\d,.]+)\s*%', s3_text)
        if bond_match_nat:
            # Find the specific line containing this match
            matched_line = ""
            for line in s3_text.splitlines():
                if bond_match_nat.group(0) in line:
                    matched_line = line
                    break
            change_val = bond_match_nat.group(1)
            change_str = change_val + "bp"
            if "하락한" in matched_line or "내린" in matched_line:
                if not any(char in change_str for char in ['△', '-', '▼']):
                    change_str = "△" + change_str
            data['macro_indicators'].append({
                'name': '국고채 금리 (10년물)',
                'value': bond_match_nat.group(2) + "%",
                'change': change_str
            })
        
    # Parse events
    s4_text = data['section4_text']
    events_found = []
    if "CPI" in s4_text or "소비자물가" in s4_text:
        events_found.append("미국 5월 소비자물가지수(CPI) 발표 예정")
    if "FOMC" in s4_text:
        events_found.append("미국 연준(Fed) FOMC 정례회의 금리 결정 예정")
    if "PPI" in s4_text or "생산자물가" in s4_text:
        events_found.append("미국 5월 생산자물가지수(PPI) 발표 예정")
    if "고용" in s4_text:
        events_found.append("미국 고용지표 변동에 따른 영향 점검")
    if "스페이스" in s4_text:
        events_found.append("스페이스 상장 등 대형 수급 일정 대기")
        
    if events_found:
        data['events'] = events_found
    else:
        data['events'] = [
            "글로벌 인플레이션 지표 발표에 따른 통화정책 기조 점검",
            "미국 연준(Fed) 금리 결정 및 통화정책 전망 분석"
        ]
        
    # Check if we have a quote in s3_text
    quote_match = re.search(r'※\s*([가-힣\s\(\)]+총재),\s*"(.*)"', s3_text)
    if quote_match:
        data['quote'] = {
            'speaker': quote_match.group(1).strip(),
            'content': quote_match.group(2).strip()
        }
    else:
        data['quote'] = {
            'speaker': '한국은행(BOK) 이창용 총재',
            'content': '지정학적 리스크 및 환율 변동성 확대 상황을 예의주시하며 시장 안정화를 위해 긴밀히 공조할 것.'
        }
    
    return data

# ==============================================================================
# DOCUMENT BUILDER
# ==============================================================================
def wrap_korean_text(text, max_len=30):
    """Wraps Korean text strictly by spaces, ensuring no line exceeds max_len characters.
    Words are never split across lines. Certain price/index/YTD lines are kept on a single line."""
    line_stripped = text.strip()
    
    # Check if this line should not be wrapped (YTD, Samsung stocks, exchange rates, bond yields)
    if line_stripped.startswith('※') or (line_stripped.startswith('*') and '연초대비' in line_stripped):
        return [text]
    if line_stripped.startswith('(') and any(kw in line_stripped for kw in ['전자', '화재', '생명']):
        return [text]
    if ('△' in line_stripped or '+' in line_stripped or '%' in line_stripped) and any(kw in line_stripped for kw in ['원', 'bp', '%', 'pt', '지수', '상승', '하락', '금리']):
        if len(line_stripped) < 75:
            return [text]
            
    is_bold_header = text.startswith('**') and text.endswith('**')
    if is_bold_header:
        text = text[2:-2].strip()
        
    words = text.split()
    lines = []
    current_line = []
    current_len = 0
    
    for word in words:
        visible_word = word.replace("**", "").replace("__", "")
        word_len = len(visible_word)
        
        extra = 1 if current_line else 0
        if current_len + word_len + extra > max_len:
            if current_line:
                lines.append(" ".join(current_line))
                current_line = [word]
                current_len = word_len
            else:
                current_line = [word]
                current_len = word_len
        else:
            current_line.append(word)
            current_len += word_len + extra
            
    if current_line:
        lines.append(" ".join(current_line))
        
    result = [l for l in lines if l]
    if is_bold_header:
        result = [f"**{l}**" for l in result]
    return result

def create_report_document(data, output_path):
    """Generates the daily market report Word document matching the sample style exactly."""
    doc = Document()
    
    # 1. Page margins (상하좌우 여백 설정: 상 1.18인치 = 85.05pt, 하/좌/우 1.0인치 = 72pt)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Pt(85.05)
        section.bottom_margin = Pt(72.0)
        section.left_margin = Pt(72.0)
        section.right_margin = Pt(72.0)
        
    # Get raw report text
    raw_text = data if isinstance(data, str) else data.get('raw_text', '')
    if not raw_text and isinstance(data, dict):
        # Fallback: reconstruct raw text from sections if raw_text key is missing
        parts = []
        sec1 = data.get('section1_text', '')
        if sec1: parts.append(sec1)
        sec2_1 = data.get('section2_text1', '')
        if sec2_1: parts.append(sec2_1)
        sec2_2 = data.get('section2_text2', '')
        if sec2_2: parts.append(sec2_2)
        sec3 = data.get('section3_text', '')
        if sec3: parts.append(sec3)
        sec4 = data.get('section4_text', '')
        if sec4: parts.append(sec4)
        raw_text = "\n\n".join(parts)

    # Process line by line
    lines = raw_text.splitlines()
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip title line at the top of the document (sample does not show it)
        line_clean = line_stripped.replace('*', '').strip()
        if line_clean.lower().startswith('title') or line_clean.lower().startswith('**title'):
            continue
        if line_stripped.startswith('**') and line_stripped.endswith('**') and '시황(' in line_stripped:
            continue
            
        # Empty line maps to empty paragraph (preserves spacing blocks)
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Wrap the line to max 30 characters
        wrapped_lines = wrap_korean_text(line_stripped, max_len=30)
        for w_line in wrapped_lines:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(8)
            p_format.line_spacing = 1.15
            
            # Enable kinsoku and wordWrap for Korean (한글 단어 잘림 방지)
            pPr = p._p.get_or_add_pPr()
            kinsoku = OxmlElement('w:kinsoku')
            kinsoku.set(qn('w:val'), '1')
            pPr.append(kinsoku)
            
            wordWrap = OxmlElement('w:wordWrap')
            wordWrap.set(qn('w:val'), '1')
            pPr.append(wordWrap)
            
            # Check if YTD line (starts with ※)
            is_ytd = w_line.startswith('※')
            color = RGBColor(0, 0, 255) if is_ytd else RGBColor(0, 0, 0)
            
            # Parse markdown bold (**...**)
            parts = re.split(r'(\*\*[^*]+\*\*)', w_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run_text = part[2:-2]
                    bold = True
                else:
                    run_text = part
                    bold = False
                    
                if run_text:
                    run = p.add_run(run_text)
                    run.font.name = '바탕체'
                    run.font.size = Pt(11)
                    run.bold = bold
                    run.font.color.rgb = color
                    
                    # XML fix for East Asian fonts rendering
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), '바탕체')
                    rFonts.set(qn('w:hAnsi'), '바탕체')
                    rFonts.set(qn('w:eastAsia'), '바탕체')
                
    doc.save(output_path)
    print(f"[성공] 워드 문서가 아시아시황 샘플 양식으로 생성되었습니다: {output_path}")

# ==============================================================================
# MAIN COMMAND LINE INTERFACE
# ==============================================================================
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Official Gaejosik Report Word Generator")
    parser.add_argument('--input', type=str, help="Path to input text file containing the raw daily report")
    parser.add_argument('--text', type=str, help="Raw daily report text string directly")
    parser.add_argument('--output', type=str, default="Daily_Market_Report.docx", help="Path to save output Word (.docx) file")
    
    args = parser.parse_args()
    
    raw_content = ""
    if args.input:
        if os.path.exists(args.input):
            with open(args.input, 'r', encoding='utf-8') as f:
                raw_content = f.read()
        else:
            print(f"[오류] 입력 파일을 찾을 수 없습니다: {args.input}")
            sys.exit(1)
    elif args.text:
        raw_content = args.text
    else:
        print("[경고] 입력 인자가 없어 디폴트 텍스트 테스트를 진행합니다.")
        raw_content = """Title : 아시아 시황(6/8)
금일 아시아 증시 시황 보고 드립니다.
아시아 증시는 미국 고용지표 호조에 따른 금리 인상 우려와 글로벌 반도체주 급락 여파로 전기전자 업종 중심으로 폭락했습니다.
※ 6.8(연초대비): 한국 △8.29%(+77.6), 일본 △4.48%(+26.4), 중국 △1.70%(△0.2)
한국 증시는 미국발 긴축 우려 및 글로벌 반도체 업황 경계심 확산으로 대형주 중심으로 폭락했으며,
특히 원달러 환율의 높은 수준 유지와 외국인의 순매도 지속이 차익실현 욕구를 자극하여
장중 유가증권시장 서킷브레이커와 매도 사이드카가 발동되는 등 전반적으로 극심한 폭락세를 보였습니다.
외국인은 오늘 코스피 시장에서 약 3,557억원을 순매도하며 21거래일째 매도 우위를 지속했고
젠슨 황 CEO 방한 모멘텀에도 불구하고 그간 기대감에 상승했던 반도체 관련 대형주들이 동반 급락했습니다.
(전자 △10.18%, 화재 △4.79%, 생명 △8.97%)
원달러는 외국인의 국내 주식 매도 우위 지속 속에 △7원(1,535원) 소폭 하락하였으나 여전히 고환율 부담을 이어갔고,
국채금리(10년)는 미 고용지표 호조에 따른 연준 긴축 장기화 우려에 +9.4bp(4.348%) 급등하며 연중 최고치를 경신했습니다.
향후에는 이번 주 예정된 미 5월 CPI 발표 및 연준의 FOMC 정례회의 등
주요 글로벌 매크로 이벤트 결과에 따라 국내외 증시의 변동성은 지속될 전망입니다."""
        
    structured_data = parse_raw_text(raw_content)
    create_report_document(structured_data, args.output)
