import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==============================================================================
# DOCX XML STYLING HELPERS
# ==============================================================================
def style_run(run, font_name='바탕체', size_pt=15, bold=False, underline=False, color_rgb=None, is_footnote=False):
    """Applies strict typography to a run in Batang font with explicit color control."""
    actual_size = 11 if is_footnote else size_pt
    actual_color = RGBColor(0, 0, 255) if is_footnote else (color_rgb if color_rgb is not None else RGBColor(0, 0, 0))
    actual_bold = False if is_footnote else bold
    actual_underline = False if is_footnote else underline
    
    run.font.name = font_name
    run.font.size = Pt(actual_size)
    run.bold = actual_bold
    run.underline = actual_underline
    run.font.color.rgb = actual_color
    
    # XML fix for East Asian font rendering in MS Word
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph(paragraph, space_before_pt, space_after_pt, left_indent_pt, first_line_indent_pt, word_wrap_off=False):
    """Formats paragraph spacing, indents, and line spacing according to guidelines."""
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before_pt)
    p_format.space_after = Pt(space_after_pt)
    p_format.left_indent = Pt(left_indent_pt)
    p_format.first_line_indent = Pt(first_line_indent_pt)
    p_format.line_spacing = 1.3  # Fixed line spacing 1.3
    
    # Use LEFT alignment to prevent wide spaces between words in Korean.
    p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Ensure word wrap does not split characters (단어가 끊어지지 않도록 줄바꿈 설정)
    if word_wrap_off:
        pPr = paragraph._p.get_or_add_pPr()
        wordWrap = OxmlElement('w:wordWrap')
        wordWrap.set(qn('w:val'), 'off')
        pPr.append(wordWrap)

# ==============================================================================
# REPORT CONTENT GENERATOR
# ==============================================================================
def main():
    doc = Document()
    
    # 1. Page margins set to 2.0 cm on all sides (상하좌우 여백 2cm 고정) & A4 Size
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
    # 2. HEADLINE (헤드라인: 22pt, 바탕체, 굵게, 밑줄)
    headline_p = doc.add_paragraph()
    headline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_p.paragraph_format.space_before = Pt(12)
    headline_p.paragraph_format.space_after = Pt(24)
    headline_p.paragraph_format.line_spacing = 1.3
    
    headline_run = headline_p.add_run("금일 아시아 및 국내 증시 시황 보고서 (6/2)")
    style_run(headline_run, font_name='바탕체', size_pt=22, bold=True, underline=True)
    
    # Data definitions (Use tabs instead of spaces after bullets)
    report_data = [
        # SECTION 1
        {
            'level1': "1.\t아시아 증시 동향 및 주요국 수익률",
            'level2': "□\t미국 증시의 사상 최고치 기록에 힘입어 아시아 주요국 증시가 동반 상승하였으나 중국은 제조업 지표 둔화로 약세를 보임.",
            'level3': [
                "-\t한국 KOSPI 지수는 전장 대비 1.08% 상승한 8,883.19로 출발하며 사상 최고치를 돌파한 후 장중 차익실현 매물 출회로 변동성이 크게 확대됨.",
                "-\t일본 Nikkei 225 지수는 전장 대비 +0.9% 상승하며 연초 대비 +33.0%의 안정적인 성장세를 기록함.",
                "-\t중국 Shanghai 종합 지수는 5월 제조업 PMI 지표의 둔화 영향으로 전장 대비 △0.2% 소폭 하락하며 연초 대비 +2.3% 수준에 머무름."
            ]
        },
        # SECTION 2
        {
            'level1': "2.\t국내 증시(KOSPI) 및 삼성그룹 주요 계열사 동향",
            'level2': "□\t코스피가 사상 최고점을 경신한 후 외국인의 대규모 매도세로 인해 장중 급격한 등락을 거듭하며 조정 국면에 진입함.",
            'level3': [
                {
                    'text': "-\t외국인 투자자가 장중 1조 5,000억 원 이상의 대규모 순매도를 기록하며 지수를 끌어내린 반면 개인과 기관은 동반 순매수로 대응함."
                },
                {
                    'text': "-\t삼성전자는 AI 반도체 및 HBM 수요 증가 기대감으로 보통주 시가총액 2,000조 원을 돌파하고 사상 최고가인 349,000원을 경신함.",
                    'level4': [
                        "·\tHBM4가 탑재된 엔비디아의 베라루빈 차세대 라인업 양산 발표가 강력한 매수세 유입의 핵심 동력으로 작용함."
                    ]
                },
                {
                    'text': "-\t삼성화재와 삼성생명은 보험업종 내 견조한 이익 체력과 주주환원 확대 기대감을 바탕으로 각각 604,000원과 410,000원의 견조한 주가 흐름을 유지함.",
                    'level4': [
                        "·\t삼성화재는 업계 최고 수준의 브랜드 가치 및 자본력을 바탕으로 증권가의 목표주가 상향 리포트가 집중됨.",
                        "·\t삼성생명은 보험서비스 손익 개선에 따른 연결 실적 성장세가 부각되며 보험 대장주로서 장기 상승 모멘텀을 확보함."
                    ]
                }
            ]
        },
        # SECTION 3
        {
            'level1': "3.\t외환(원달러) 및 채권 시장 동향",
            'level2': "□\t외국인의 지속적인 주식 순매도세와 매파적인 한국은행(BOK) 인사 발언의 영향으로 환율과 국채금리가 동반 급등함.",
            'level3': [
                {
                    'text': "-\t원달러 환율은 서울 외환시장에서 전 거래일 대비 9.40원 급등한 1,513.70원에 주간거래를 시작하며 고환율 불안 기조를 이어감."
                },
                {
                    'text': "-\t국고채 10년물 금리는 BOK 국제콘퍼런스에서의 매파적 통화정책 발언 여파로 전장 대비 +10bp 상승한 4.17%를 기록함.",
                    'level4': [
                        "·\t※ 신현송 총재는 단기적인 인플레이션 압력이 이미 높은 수준이며 이에 대응하기 위한 정책적 운신의 폭이 확대 되었음을 재차 강조함."
                    ]
                }
            ]
        },
        # SECTION 4
        {
            'level1': "4.\t금주 주요 일정 및 투자 전망",
            'level2': "□\t엔비디아 GTC 대만 및 컴퓨텍스 2026 등 AI 반도체 관련 대형 이벤트의 집중으로 관련 수혜주로의 자금 쏠림과 변동성 확대에 유의함.",
            'level3': [
                {
                    'text': "-\t엔비디아 GTC 대만(1~4일) 및 대만 컴퓨텍스 2026(2~5일) 개최에 따라 글로벌 테크 기업들의 신기술 공개 일정이 집중됨."
                },
                {
                    'text': "-\t젠슨 황 엔비디아 CEO의 6월 8일 방한이 유력해짐에 따라 HBM 관련 밸류체인 종목들의 단기 변동성 확대 우려가 상존함.",
                    'level4': [
                        "·\t대형 주도주 중심의 수급 쏠림 현상이 단기적으로 심화될 가능성이 높아 포트폴리오 다변화 관점에서의 접근이 권장됨."
                    ]
                }
            ]
        }
    ]
    
    # 3. Add Content blocks
    for section in report_data:
        # --- Level 1 ("1.") ---
        p1 = doc.add_paragraph()
        format_paragraph(p1, space_before_pt=12, space_after_pt=12, left_indent_pt=22.5, first_line_indent_pt=-22.5)
        run1 = p1.add_run(section['level1'])
        style_run(run1, font_name='바탕체', size_pt=15, bold=True)
        
        # --- Level 2 ("□") ---
        p2 = doc.add_paragraph()
        format_paragraph(p2, space_before_pt=12, space_after_pt=12, left_indent_pt=22.5, first_line_indent_pt=-22.5)
        run2 = p2.add_run(section['level2'])
        style_run(run2, font_name='바탕체', size_pt=15, bold=True)
        
        # --- Level 3 ("-") ---
        for l3_item in section['level3']:
            if isinstance(l3_item, dict):
                text_content = l3_item['text']
                level4_list = l3_item.get('level4', [])
            else:
                text_content = l3_item
                level4_list = []
                
            is_fn3 = "※" in text_content
            p3 = doc.add_paragraph()
            format_paragraph(p3, space_before_pt=6, space_after_pt=6, left_indent_pt=30.0, first_line_indent_pt=-15.0, word_wrap_off=True)
            
            run3 = p3.add_run(text_content)
            style_run(run3, font_name='바탕체', size_pt=15, bold=False, is_footnote=is_fn3)
            
            # --- Level 4 ("·") ---
            for l4_text in level4_list:
                is_fn4 = "※" in l4_text
                p4 = doc.add_paragraph()
                format_paragraph(p4, space_before_pt=3, space_after_pt=3, left_indent_pt=37.5, first_line_indent_pt=-15.0, word_wrap_off=True)
                
                run4 = p4.add_run(l4_text)
                style_run(run4, font_name='바탕체', size_pt=15, bold=False, is_footnote=is_fn4)
                
    output_path = "Daily_Market_Report_Official_20260602.docx"
    doc.save(output_path)
    print(f"[성공] 공식 공공기관 보고서 형식 워드 문서 생성 완료: {output_path}")

if __name__ == '__main__':
    main()
